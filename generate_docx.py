import sys
from docx import Document
from docx.shared import Inches, Pt
import os

# Usage: python3 generate_docx.py input.txt logo.png
input_txt = sys.argv[1]
logo_path = sys.argv[2]
output_docx = "/tmp/SmartEmail_Output.docx"

doc = Document()

# Insert logo if available
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2))
    doc.add_paragraph("")

# Heading
heading = doc.add_paragraph("SmartEmail Response")
heading.runs[0].bold = True
heading.runs[0].font.size = Pt(16)
doc.add_paragraph("")

# Insert text content
with open(input_txt, "r", encoding="utf-8") as f:
    text = f.read().strip()

# Split sections if marked (e.g. for Generate and Enhance)
if "===ENHANCED===" in text:
    original, enhanced = text.split("===ENHANCED===")
    doc.add_heading("Original Message:", level=2)
    doc.add_paragraph(original.strip())
    doc.add_paragraph("")
    doc.add_heading("Enhanced Version:", level=2)
    doc.add_paragraph(enhanced.strip())
else:
    doc.add_paragraph(text)

doc.save(output_docx)
